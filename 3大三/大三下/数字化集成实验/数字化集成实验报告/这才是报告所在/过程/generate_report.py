from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
TEMPLATE_PATH = ROOT / "模板" / "《数字化集成控制系统》期末报告模板.docx"
OUTPUT_PATH = ROOT / "数字化集成控制系统期末实验报告.docx"
ASSET_DIR = ROOT / "generated_assets"
ASSET_DIR.mkdir(exist_ok=True)

CODE_FILES = {
    "附录A 定位跑 STL 程序": ROOT / "STL代码" / "定位跑.txt",
    "附录B 折返跑 STL 程序": ROOT / "STL代码" / "折返跑.txt",
    "附录C 跨栏跑 STL 程序": ROOT / "STL代码" / "跨栏跑.txt",
    "附录D 位置 PID STL 程序": ROOT / "STL代码" / "位置PID.txt",
    "附录E 位置/速度双闭环 PID STL 程序": ROOT / "STL代码" / "双闭环PID.txt",
}


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = []
    if bold:
        candidates.extend(
            [
                r"C:\Windows\Fonts\msyhbd.ttc",
                r"C:\Windows\Fonts\simhei.ttf",
            ]
        )
    candidates.extend(
        [
            r"C:\Windows\Fonts\msyh.ttc",
            r"C:\Windows\Fonts\simsun.ttc",
            r"C:\Windows\Fonts\simsun.ttc",
        ]
    )
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def set_run_font(run, size: float, bold: bool = False, font_name: str = "宋体") -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    r_fonts = run._element.rPr.rFonts
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)


def clear_paragraph(paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def delete_paragraph(paragraph) -> None:
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def configure_paragraph(paragraph, first_line_indent_pt: float = 22, line_spacing: float = 1.5) -> None:
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(first_line_indent_pt) if first_line_indent_pt else Pt(0)
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def add_text_paragraph(
    doc: Document,
    text: str,
    *,
    size: float = 11,
    bold: bool = False,
    first_line_indent_pt: float = 22,
    line_spacing: float = 1.5,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = alignment
    configure_paragraph(paragraph, first_line_indent_pt, line_spacing)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold=bold)


def add_heading(doc: Document, text: str, size: float, *, bold: bool = True) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    configure_paragraph(paragraph, first_line_indent_pt=0, line_spacing=1.2)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold=bold)


def add_figure(doc: Document, image_path: Path, caption: str, width_cm: float = 15.5) -> None:
    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    configure_paragraph(pic_para, first_line_indent_pt=0, line_spacing=1.0)
    pic_para.add_run().add_picture(str(image_path), width=Cm(width_cm))

    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    configure_paragraph(cap_para, first_line_indent_pt=0, line_spacing=1.0)
    run = cap_para.add_run(caption)
    set_run_font(run, 10.5)


def wrap_text_by_chars(text: str, max_chars: int) -> list[str]:
    lines: list[str] = []
    buffer = ""
    for char in text:
        buffer += char
        if len(buffer) >= max_chars:
            lines.append(buffer)
            buffer = ""
    if buffer:
        lines.append(buffer)
    return lines or [text]


def draw_centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill: str) -> None:
    lines = wrap_text_by_chars(text, 12)
    line_boxes = [draw.textbbox((0, 0), line, font=font) for line in lines]
    total_height = sum(b[3] - b[1] for b in line_boxes) + 8 * (len(lines) - 1)
    x1, y1, x2, y2 = box
    current_y = y1 + ((y2 - y1) - total_height) / 2
    for line, bbox in zip(lines, line_boxes):
        width = bbox[2] - bbox[0]
        height = bbox[3] - bbox[1]
        current_x = x1 + ((x2 - x1) - width) / 2
        draw.text((current_x, current_y), line, font=font, fill=fill)
        current_y += height + 8


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], *, fill: str = "#384860", width: int = 5) -> None:
    draw.line([start, end], fill=fill, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        arrow = [(ex, ey), (ex - 16 * direction, ey - 8), (ex - 16 * direction, ey + 8)]
    else:
        direction = 1 if ey > sy else -1
        arrow = [(ex, ey), (ex - 8, ey - 16 * direction), (ex + 8, ey - 16 * direction)]
    draw.polygon(arrow, fill=fill)


def draw_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fill: str, outline: str, font) -> None:
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=4)
    draw_centered_text(draw, box, text, font, "#1C2430")


def create_hardware_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 1100), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34, bold=True)
    box_font = load_font(26)
    small_font = load_font(22)

    draw.text((560, 40), "数字化集成控制系统硬件拓扑图", font=title_font, fill="#1B2B45")

    boxes = {
        "power": (90, 170, 430, 310),
        "hmi": (120, 470, 460, 610),
        "plc": (620, 310, 1030, 470),
        "encoder": (1180, 120, 1580, 260),
        "vfd": (1180, 360, 1580, 520),
        "motor": (1180, 680, 1580, 840),
    }

    draw_box(draw, boxes["power"], "电源模块", "#DDEBF7", "#5B8DB8", box_font)
    draw_box(draw, boxes["hmi"], "HMI 触摸屏\n数值输入/状态显示", "#E2F0D9", "#7FA66A", box_font)
    draw_box(draw, boxes["plc"], "S7-200 PLC\n高速计数/逻辑控制/PID运算", "#FCE4D6", "#D48758", box_font)
    draw_box(draw, boxes["encoder"], "增量式编码器\nA/B 相位置反馈", "#FFF2CC", "#C7A64C", box_font)
    draw_box(draw, boxes["vfd"], "FRS520SE 变频器\n接收方向与速度指令", "#E4DFEC", "#8A76AF", box_font)
    draw_box(draw, boxes["motor"], "电机执行机构\n实现定位、折返、跨栏运动", "#F4CCCC", "#B15D5D", box_font)

    draw_arrow(draw, (460, 540), (620, 390))
    draw.text((470, 500), "RS485 通讯", font=small_font, fill="#384860")

    draw_arrow(draw, (1180, 190), (1030, 360))
    draw.text((1080, 215), "I0.0/I0.1", font=small_font, fill="#384860")

    draw_arrow(draw, (1030, 410), (1180, 440))
    draw.text((1010, 455), "Q0.0/Q0.1\nAQW0/Q0.3~Q0.5", font=small_font, fill="#384860")

    draw_arrow(draw, (1380, 520), (1380, 680))
    draw.text((1415, 580), "驱动输出", font=small_font, fill="#384860")

    draw_arrow(draw, (430, 240), (620, 360))
    draw_arrow(draw, (360, 310), (360, 470))
    draw_arrow(draw, (430, 240), (1180, 440))
    draw.text((450, 205), "供电与接口电源", font=small_font, fill="#384860")

    image.save(path)


def create_software_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 1300), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34, bold=True)
    box_font = load_font(24)
    small_font = load_font(20)

    draw.text((540, 40), "主程序与子程序软件流程图", font=title_font, fill="#1B2B45")

    steps = [
        ("系统上电初始化\n配置 HSC0、定时中断、PID 参数", (620, 120, 1180, 250), "#DDEBF7", "#5B8DB8"),
        ("读取 HMI 参数\n圈数、折返点、速度切换点", (620, 310, 1180, 440), "#E2F0D9", "#7FA66A"),
        ("模式判断\n定位跑 / 折返跑 / 跨栏跑 / PID 模式", (620, 500, 1180, 630), "#FCE4D6", "#D48758"),
        ("执行对应子程序\n状态机/计数器/PID 运算", (620, 690, 1180, 820), "#E4DFEC", "#8A76AF"),
        ("输出控制量\n方向输出、速度档位、模拟量", (620, 880, 1180, 1010), "#FFF2CC", "#C7A64C"),
        ("编码器反馈与结束判断\n到位、折返、变速、停机", (620, 1070, 1180, 1200), "#F4CCCC", "#B15D5D"),
    ]

    for text, box, fill, outline in steps:
        draw_box(draw, box, text, fill, outline, box_font)

    for idx in range(len(steps) - 1):
        start_box = steps[idx][1]
        end_box = steps[idx + 1][1]
        draw_arrow(draw, ((start_box[0] + start_box[2]) // 2, start_box[3]), ((end_box[0] + end_box[2]) // 2, end_box[1]))

    draw_arrow(draw, (1180, 1135), (1400, 1135))
    draw_arrow(draw, (1400, 1135), (1400, 375))
    draw_arrow(draw, (1400, 375), (1180, 375))
    draw.text((1425, 720), "未结束则循环\n实时刷新参数", font=small_font, fill="#384860")

    draw_box(draw, (120, 520, 460, 690), "高速计数子程序\nHC0/脉冲换算", "#D9EAD3", "#6D9E6A", box_font)
    draw_box(draw, (120, 760, 460, 930), "PID 运算子程序\n位置环/速度环", "#D0E0E3", "#5E8CA5", box_font)
    draw_arrow(draw, (460, 605), (620, 565))
    draw_arrow(draw, (460, 845), (620, 755))
    draw.text((165, 950), "子程序结果写回主流程", font=small_font, fill="#384860")

    image.save(path)


def prepare_template(doc: Document) -> None:
    title_para = doc.paragraphs[0]
    clear_paragraph(title_para)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    configure_paragraph(title_para, first_line_indent_pt=0, line_spacing=1.0)
    title_run = title_para.add_run("《数字化集成控制系统》期末实验报告")
    set_run_font(title_run, 24, bold=True)

    # Delete all instructional paragraphs and the final topic image.
    for paragraph in list(doc.paragraphs[26:])[::-1]:
        delete_paragraph(paragraph)


def append_report_body(doc: Document) -> None:
    add_heading(doc, "一、摘要与关键词", 16)
    add_text_paragraph(
        doc,
        "摘要：本报告基于数字化集成控制系统课程实验素材，围绕定位跑、折返跑、跨栏跑、位置 PID 定位跑以及位置/速度双闭环 PID 折返跑五个实验题展开。系统以 S7-200 PLC 为核心，结合编码器反馈、HMI 参数输入与变频器驱动，实现了从开环顺序控制到闭环精确调速定位的逐步升级。报告首先给出系统总体需求、硬件连接关系和软件流程结构，然后说明高速计数、闭环控制、PID 调节等基础原理，最后结合 STL 程序分析各实验题的控制逻辑、关键变量、运行流程与改进效果。整体实现表明，该系统能够满足课程要求中的多模式运动控制任务，并为后续完成更高精度的位置控制和参数在线整定打下基础。",
        first_line_indent_pt=0,
    )
    add_text_paragraph(
        doc,
        "关键词：S7-200 PLC；高速计数；定位控制；PID；数字化集成控制系统",
        first_line_indent_pt=0,
    )

    add_heading(doc, "1 系统总体方案设计", 16)
    add_heading(doc, "1.1 系统需求分析", 15)
    add_text_paragraph(
        doc,
        "本次课程实验要求围绕同一套电机执行机构完成 5 个不同层级的运动控制任务。基础层面需要实现定圈定位跑、带折返点的折返跑以及带速度切换点的跨栏跑；进阶层面需要引入编码器反馈和 PID 算法，实现位置闭环定位以及位置/速度双闭环折返控制。由此可知，系统不仅要具备可靠的顺序控制能力，还必须具备高速脉冲采集、状态切换、运行停止判定和在线参数调整能力。",
    )
    add_text_paragraph(
        doc,
        "结合题目要求与现有程序，系统的核心需求可归纳为四点：一是运动需求，即在不同模式下完成定距停止、往返切换和分段变速；二是交互需求，即通过 HMI 数值输入控件对圈数、折返点和速度切换点进行调整；三是控制需求，即在 PID 模式下利用编码器脉冲实现闭环修正，抑制超调和振荡；四是工程需求，即保证 PLC 程序结构清晰、便于调试和扩展。现有 AI 聊天记录主要用于辅助梳理 HSC0 配置、状态机切换思路和超调修正策略，正文仅保留这些已经验证有效的成果。"
    )
    add_text_paragraph(
        doc,
        "从实验装置和程序参数可看出，系统默认以编码器脉冲作为位置量，利用 Q0.0/Q0.1 完成方向控制，利用 Q0.3~Q0.5 或 AQW0 完成速度输出。基于课程任务，报告采用“先实现基础顺序控制，再引入闭环反馈和 PID 算法”的总体设计思路，使五个实验题之间形成逐层递进的关系。"
    )

    add_heading(doc, "1.2 系统硬件与软件架构设计", 15)
    add_heading(doc, "1.2.1 硬件架构设计", 11, bold=False)

    hardware_path = ASSET_DIR / "hardware_topology.png"
    software_path = ASSET_DIR / "software_flow.png"
    create_hardware_diagram(hardware_path)
    create_software_diagram(software_path)

    add_text_paragraph(
        doc,
        "系统硬件由 PLC 控制器、增量式编码器、HMI 触摸屏、变频器、电机以及电源模块组成。PLC 是整个系统的控制核心，负责读取编码器反馈、解析 HMI 输入参数、执行顺序逻辑和 PID 运算；编码器提供位置脉冲，接入 PLC 的高速计数输入；变频器负责接收 PLC 的方向与速度指令并驱动电机；HMI 则承担参数输入、状态显示和运行监视等任务。整体结构形成“参数设定—控制运算—执行输出—位置反馈”的闭环链路。",
    )
    add_figure(doc, hardware_path, "图1 系统硬件拓扑图")
    add_text_paragraph(
        doc,
        "根据现有代码与说明书可知，本系统重点使用了 S7-200 的高速计数能力以及模拟量/开关量输出能力。编码器 A/B 相信号接入 I0.0、I0.1，用于获取电机实时位置；Q0.0、Q0.1 输出正反转控制；Q0.3、Q0.4、Q0.5 表示不同速度档位；在位置 PID 实验中，AQW0 进一步被用于向变频器输出 0~10V 模拟量，实现连续速度调节。"
    )

    add_heading(doc, "1.2.2 软件架构设计", 11, bold=False)
    add_text_paragraph(
        doc,
        "软件结构采用主程序加功能子程序的分层设计。主程序负责系统初始化、模式判断、参数装载和结束判定；功能层则按照实验任务分别实现定位跑、折返跑、跨栏跑、位置 PID 和双闭环 PID 的具体控制逻辑。这样做的优点是各个实验题能够共享同一套输入输出资源与基础变量，同时又可以在算法层面保持相对独立，便于逐题调试与性能比较。",
    )
    add_figure(doc, software_path, "图2 主程序与子程序软件流程图")
    add_text_paragraph(
        doc,
        "在运行流程上，程序首先完成 HSC0 与中断参数初始化，然后读取 HMI 或实验预设参数，接着根据当前模式调用对应的控制程序。执行过程中，编码器反馈不断回送至 PLC，用于位置换算、折返点判断或者 PID 过程变量更新；当系统检测到到位、到圈数或状态切换条件满足时，程序自动停止或进入下一阶段。这种结构既能支持基础的顺序控制，也能兼容进阶的闭环控制。"
    )

    add_heading(doc, "2 基础控制原理与概念", 16)
    add_text_paragraph(
        doc,
        "本课程实验涉及电机位置控制、速度调节、编码器反馈、高速计数与 PID 调节等多项基础原理。理解这些概念，有助于说明为什么前 3 个实验更偏向顺序逻辑和计数判定，而后 2 个实验则能够实现更高精度、更平滑的闭环控制。",
    )

    add_heading(doc, "2.1 电机运动控制基础概念", 15)
    add_heading(doc, "2.1.1 定位控制与闭环控制", 11, bold=False)
    add_text_paragraph(
        doc,
        "定位控制的目标是在给定距离或给定脉冲数处停止电机。最简单的实现方式是基于计数阈值进行开关控制，即当累计脉冲达到目标值时切断输出；这种方式结构简单，适用于定位跑等基础任务。闭环控制则进一步把“目标位置”和“实际位置”进行比较，根据偏差持续修正输出，因此更适合需要减小超调、降低振荡和提高终点稳定性的场景。",
    )

    add_heading(doc, "2.1.2 增量式编码器工作原理", 11, bold=False)
    add_text_paragraph(
        doc,
        "增量式编码器通过 A/B 两相信号输出脉冲，PLC 可根据脉冲个数判断电机位移，根据相位关系判断转动方向。实验中，编码器的脉冲数被用作位置量或测速量。对于基础实验，累计脉冲达到设定值即可触发停车或变速；对于 PID 实验，编码器脉冲还会被换算为标准化过程变量，为控制器提供实时反馈。",
    )

    add_heading(doc, "2.1.3 PLC高速计数器与HMI交互原理", 11, bold=False)
    add_text_paragraph(
        doc,
        "S7-200 的高速计数器能够在普通扫描周期之外对高速脉冲进行可靠采集，这对于电机位置控制十分关键。现有定位跑和 PID 程序均配置了 HSC0 模式，通过 SMB37、SMD38 等寄存器完成计数器定义、清零和当前值更新。另一方面，HMI 提供数值输入控件，可用于在线设置圈数、折返点和变速位置，使系统从固定参数运行逐步过渡到可调参数运行，增强了实验系统的人机交互能力。"
    )

    add_heading(doc, "2.2 PID控制理论基础", 15)
    add_text_paragraph(
        doc,
        "PID 控制由比例 P、积分 I、微分 D 三部分组成。比例环节根据当前偏差快速给出响应，决定系统的跟随强度；积分环节用于消除静差，使系统更接近目标值；微分环节反映偏差变化趋势，有助于抑制超调和振荡。在位置 PID 实验中，程序采用标准化的位置偏差驱动模拟量输出，通过“接近终点时切断动力”的死区策略实现平稳停车；在双闭环 PID 实验中，外环根据位置误差生成目标速度，内环再根据速度误差调节控制输出，从而兼顾响应速度与最终定位精度。"
    )

    add_heading(doc, "3 控制方案设计", 16)
    add_text_paragraph(
        doc,
        "本节结合五份 STL 程序说明各实验题的实现思路、关键变量和控制要点。为了避免正文篇幅过大，完整程序统一整理于附录，正文仅给出必要的结构分析和实验结果说明。",
    )

    add_heading(doc, "3.1 定位跑控制方案与理论要点", 15)
    add_text_paragraph(
        doc,
        "定位跑程序的目标是在电机运行 3 圈后自动停车。程序在首次扫描时初始化 HSC0，并将目标值设为 1200 脉冲；启动时通过 M10.0、M10.3、M10.4 实现运行标志与方向锁存，通过清零 SMD38 保证每次运行都从统一零点开始。运行过程中，PLC 持续将 HC0 当前值搬移到 VD100，当 VD100 大于等于目标值时置位 M10.5，随后复位运行和方向输出，实现自动停车。完整 STL 程序见附录A。",
    )
    add_text_paragraph(
        doc,
        "该方案的优点是结构简洁、实现快速，适合课程实验中的基础定位任务。由于采用“到阈值即停机”的方式，程序逻辑清晰，便于验证编码器反馈和高速计数器配置是否正确；不足之处在于停车动作更接近开关控制，对速度和惯性变化的适应性有限，因此在后续 PID 实验中需要进一步引入闭环调节。"
    )

    add_heading(doc, "3.2 折返跑控制方案与理论要点", 15)
    add_text_paragraph(
        doc,
        "折返跑程序主要采用计数器与状态分段相结合的方式实现往返运动。程序先利用 C0 对编码器脉冲进行累积，达到一圈后触发中间状态，再借助 C1、C2 对后续区段进行判定，通过对 Q0.0、Q0.1 方向输出和 Q0.3 速度输出的置位/复位实现运动方向切换与阶段运行。该思路本质上是把整个往返过程划分为多个带条件的运动段，达到折返点后切换控制对象。完整 STL 程序见附录B。",
    )
    add_text_paragraph(
        doc,
        "从实验意义上看，折返跑是由单一终点定位向多阶段顺序控制的过渡。它要求程序不仅能记录位置，还要根据不同计数区间决定后续动作，因此对状态切换的可靠性提出了更高要求。该实验为后续双闭环 PID 折返跑中的状态机设计提供了基础经验。"
    )

    add_heading(doc, "3.3 跨栏跑控制方案与理论要点", 15)
    add_text_paragraph(
        doc,
        "跨栏跑程序面向“全程 7 圈、分段变速”的任务要求。程序使用 C0 记录圈数脉冲，并进一步通过 C1、C2、C3、C4 构造多个位置节点，利用 M0.0 作为总运行标志控制 Q0.0 正转输出，同时根据不同区间置位 Q0.2、Q0.3、Q0.4 三个速度档。这样就可以在起始段、中间段和后段分别采用不同的转速，实现类似跨栏的节奏变化控制。完整 STL 程序见附录C。",
    )
    add_text_paragraph(
        doc,
        "该实验的理论重点在于“位置触发变速”。与定位跑仅判断终点不同，跨栏跑要求程序在途中多个位置点切换速度，因此控制逻辑从单点停止扩展为多点事件响应。现有程序已经建立了多个圈数节点，并预留了进一步扩展三个中间位置可配置输入的接口思路，说明系统具备从固定变速向参数化变速继续升级的可能。"
    )

    add_heading(doc, "3.4 其他选做题目", 15)
    add_text_paragraph(
        doc,
        "除前三个基础实验外，本次实验还完成了两个更能体现控制深度的选做题：位置 PID 定位跑，以及位置/速度双闭环 PID 折返跑。它们都利用编码器反馈构成闭环控制，并且在程序结构上明显比基础顺序控制更复杂。",
    )

    add_heading(doc, "3.4.1 位置 PID 定位跑", 11, bold=False)
    add_text_paragraph(
        doc,
        "位置 PID 程序在上电时完成 HSC0 配置与 PID 参数初始化，将目标位置设为 1185 脉冲，并把实际脉冲数除以 2000 后标准化为 0.0~1.0 的过程变量。PLC 调用原生 PID 指令后，将输出量 VD108 换算为模拟量 AQW0，对变频器施加连续速度控制；当脉冲达到 1184 左右时，程序立即切断 Q0.0 并把 AQW0 拉低至 0V，以形成稳定的刹车死区。完整 STL 程序见附录D。",
    )
    add_text_paragraph(
        doc,
        "与基础定位跑相比，位置 PID 模式能够在接近终点时提前减速，显著减小因惯性带来的越位和抖动。现有 AI 对话文件中与“位置 PID 控制实现”“调节正转超调”等主题相关的记录，也表明该部分工作的有效成果主要集中在目标归一化、输出限幅和超调抑制策略上。"
    )

    add_heading(doc, "3.4.2 位置/速度双闭环 PID 折返跑", 11, bold=False)
    add_text_paragraph(
        doc,
        "双闭环 PID 程序在首次扫描时初始化 HSC0、50ms 定时中断、位置历史量和速度环积分器，并通过 VW100 构造完整的状态机。外部启动后，程序先把当前脉冲清零并进入状态1，目标位置设为 1200 脉冲；达到第一终点后切换到状态2，目标回到 0；随后再进入状态3，重新向 1200 脉冲前进，最终以状态4 表示任务完成。整个过程中，位置外环根据偏差生成目标速度，速度内环再根据测速结果进行 PI 调节。完整 STL 程序见附录E。",
    )
    add_text_paragraph(
        doc,
        "双闭环结构的优势在于把“跑多远”和“跑多快”两个问题分开处理。位置外环决定系统应以多大速度逼近终点，速度内环负责把实际速度稳定到目标值，因此相比单环位置控制更容易兼顾响应速度、运行平稳性和重复定位精度。这也是从课程实验走向工程控制的重要一步。"
    )

    add_heading(doc, "4 总结与展望", 16)
    add_text_paragraph(
        doc,
        "通过本次数字化集成控制系统实验，可以清楚看到电机控制任务从基础顺序逻辑到闭环算法控制的递进过程。定位跑验证了高速计数器与方向输出的基本用法；折返跑和跨栏跑进一步引入多阶段位置判断与变速逻辑；位置 PID 与双闭环 PID 则在编码器反馈基础上实现了更加平滑和精确的控制效果。五个实验题共同构成了一个较完整的电机运动控制训练链条。",
    )
    add_text_paragraph(
        doc,
        "在实现过程中，AI 聊天记录主要承担了辅助作用，例如帮助梳理 HSC0 初始化步骤、修正顺序逻辑、分析 PID 参数对超调的影响等；最终保留下来的有效成果已经固化为现有 STL 程序和本报告中的控制结论。后续若继续完善系统，可进一步把定位圈数、折返点和多段速度切换点全部开放给 HMI 输入，并增加报警、故障自检和参数整定界面，使课程实验系统更接近工程应用。"
    )

    add_heading(doc, "参考文献", 16)
    references = [
        "[1] 西门子（中国）有限公司. S7-200可编程控制器系统手册[Z]. 2005.",
        "[2] Siemens AG. TP177A/TP177B/OP177B Operating Instructions[Z].",
        "[3] FRS520SE变频器说明书[Z].",
        "[4] 胡寿松. 自动控制原理[M]. 北京: 科学出版社.",
        "[5] 廖常初. PLC编程及应用[M]. 北京: 机械工业出版社.",
    ]
    for ref in references:
        add_text_paragraph(doc, ref, first_line_indent_pt=0)

    doc.add_page_break()
    add_heading(doc, "附录 STL 程序清单", 16)
    add_text_paragraph(
        doc,
        "附录给出本次实验使用的 5 份 STL 程序。为了控制篇幅并保持代码可读性，以下内容统一采用等宽字体、小字号和紧凑行距排版。",
        first_line_indent_pt=0,
    )

    for appendix_title, file_path in CODE_FILES.items():
        add_heading(doc, appendix_title, 15)
        code_text = file_path.read_text(encoding="utf-8")
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        configure_paragraph(paragraph, first_line_indent_pt=0, line_spacing=1.0)
        paragraph.paragraph_format.left_indent = Pt(8)
        run = paragraph.add_run(code_text)
        set_run_font(run, 8.5, font_name="Courier New")


def build_report() -> None:
    doc = Document(str(TEMPLATE_PATH))
    prepare_template(doc)
    append_report_body(doc)
    doc.save(str(OUTPUT_PATH))


if __name__ == "__main__":
    build_report()
