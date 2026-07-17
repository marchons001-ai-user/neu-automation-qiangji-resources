from __future__ import annotations

import shutil
import subprocess
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
ASSET_DIR = ROOT / "dual_report_assets"
BUILD_DIR = ROOT / "dual_report_build"
ASSET_DIR.mkdir(exist_ok=True)
BUILD_DIR.mkdir(exist_ok=True)

STUDENT_CLASS = "强基班"
STUDENT_NAME = "匿名"
STUDENT_ID = ""
COURSE_NAME = "数字化集成控制系统"


IMAGE_MAP = {
    "hardware_topology.png": ROOT / "latex_assets" / "hardware_topology.png",
    "software_flow.png": ROOT / "latex_assets" / "software_flow.png",
    "hmi_runtime.png": ROOT / "屏幕截图 2026-04-24 140916.png",
    "hmi_connection.png": ROOT / "屏幕截图 2026-04-24 141103.png",
    "hmi_variables.png": ROOT / "屏幕截图 2026-04-24 141125.png",
    "plc_logic_page1.png": ROOT / "屏幕截图 2026-04-24 141333.png",
    "plc_logic_page2.png": ROOT / "屏幕截图 2026-04-24 141359.png",
}


@dataclass
class Figure:
    image: str
    caption: str
    width: float = 15.5


def copy_assets() -> None:
    for target_name, src in IMAGE_MAP.items():
        shutil.copyfile(src, ASSET_DIR / target_name)


def tex_escape(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def set_run_font(run, size: float, *, bold: bool = False, font_name: str = "宋体") -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)


def add_paragraph(
    doc: Document,
    text: str,
    *,
    size: float = 12,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent: float = 21,
    line_spacing: float = 1.35,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Pt(first_line_indent) if first_line_indent else Pt(0)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size, bold=bold)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def theory_sections() -> list[tuple[str, object]]:
    return [
        ("abstract", "本报告在前期数字化集成控制系统实验报告的基础上，补充了 TP177B 触摸屏组态、变量映射和 PLC 联调过程的截图材料，形成一份偏理论分析的综合报告。全文围绕系统总体结构、HMI 与 PLC 之间的通信机理、变量地址映射关系、操作界面的人机交互逻辑以及 PLC 程序中的互锁控制展开，重点说明触摸屏监控系统在定位跑、折返跑、跨栏跑及后续闭环控制实验中的作用。新增截图表明，本实验已完成以太网连接建立、变量表配置、触摸屏主界面设计以及主程序关键网络下载验证，从而使原有运动控制系统具备了更完整的监控、输入与显示能力。"),
        ("keywords", "关键词：S7-200 PLC；TP177B 触摸屏；WinCC flexible；变量映射；人机交互"),
        ("heading1", "1 系统功能需求与总体结构"),
        ("para", "本实验系统仍以 S7-200 PLC 为核心，完成定位跑、折返跑、跨栏跑和后续 PID 控制任务。与前版报告相比，本次补充的重点是触摸屏监控层：通过 TP177B 触摸屏建立与 PLC 的以太网连接，实现高速计数值、模拟量值和运行状态的可视化显示，同时把部分模式信号映射到人机界面，为实验参数观察、运行切换和调试记录提供直观支撑。"),
        ("para", "从系统层次上看，PLC 负责底层逻辑判断与输出执行，编码器负责位置反馈，变频器负责速度与方向驱动，而触摸屏位于监控层，承担数据显示、按键输入和状态反馈功能。也就是说，前版报告中的控制主体结构保持不变，新截图证明监控与交互层已经补齐，系统因此从“能运行”提升到了“可观察、可切换、可调试”的阶段。"),
        ("figure", Figure("hardware_topology.png", "图1 系统硬件拓扑图", 15.0)),
        ("figure", Figure("software_flow.png", "图2 主程序与子程序软件流程图", 15.0)),
        ("heading1", "2 触摸屏监控系统的理论设计"),
        ("heading2", "2.1 HMI 与 PLC 的通信机理"),
        ("para", "从连接组态截图可以看到，触摸屏型号为 TP177B 6'' color PN/DP，通信驱动程序选择为 SIMATIC S7 200，连接名为“连接_1”，接口方式为以太网。组态中 HMI 设备地址设置为 192.168.1.2，PLC 设备地址设置为 192.168.1.4，形成同一网段内的点对点工业以太网通信。这种设置满足实验环境对组态简洁性和稳定性的要求，也便于在 WinCC flexible 与 STEP7-Micro/WIN 之间进行联调。"),
        ("figure", Figure("hmi_connection.png", "图3 触摸屏与 PLC 的连接组态界面", 15.8)),
        ("heading2", "2.2 变量映射与数据一致性"),
        ("para", "变量表截图表明，HMI 已经建立了多项与 PLC 地址直接绑定的变量。例如，DWord 类型变量绑定到 VD22，用于显示高速计数或累计数值；Int 类型变量绑定到 VW20，用于显示模拟量或中间运算结果；Bool 类型变量则分别映射到 M0.2、M0.3、M0.4 和 Q1.0、Q1.1 等内部继电器与输出位。这种“数据类型—PLC 地址—采集周期”的一一对应关系，是人机界面能够稳定显示底层状态的理论基础。"),
        ("para", "从控制理论角度看，变量映射不仅是显示问题，更关系到信息一致性。若数据类型、字长或刷新周期配置错误，则触摸屏上显示的高速计数值、模式状态和模拟量反馈可能与 PLC 实际内部数据不一致，从而影响实验结论。此次变量表中统一设置了 1 s 的采集周期，说明本系统更强调稳定显示和调试观察，而非毫秒级高速刷新。"),
        ("figure", Figure("hmi_variables.png", "图4 触摸屏变量表与 PLC 地址映射", 15.8)),
        ("heading2", "2.3 界面设计与人机交互逻辑"),
        ("para", "运行界面截图显示，触摸屏主界面设置了三个明显的功能按钮区域，并提供高速计数器值、模拟量值以及起始/启动操作区。界面中部的输入控件与下方显示框形成了“输入—运行—反馈”的闭环交互结构，右下角则设置了页面返回按钮。理论上，这样的布局能够把操作按钮、数据显示和页面导航区分开来，降低误操作概率，也更符合课程实验中“先观察、后切换、再判断”的操作流程。"),
        ("figure", Figure("hmi_runtime.png", "图5 触摸屏运行界面效果图", 12.5)),
        ("heading1", "3 PLC 主程序与触摸屏联动机理"),
        ("heading2", "3.1 通信初始化与状态准备"),
        ("para", "PLC 梯形图截图显示，程序中单独设置了 ETH1_CTRL 通信相关网络，并利用 SM0.0 等系统位在上电时进行通信初始化。该网络的作用在于保证 PLC 以太网通信功能处于可用状态，为触摸屏读取内部状态位和数据寄存器提供基础。若通信初始化失败，则后续 HMI 的变量刷新与按钮信号读取都会失去依据。"),
        ("figure", Figure("plc_logic_page1.png", "图6 PLC 程序中的通信初始化与前段网络", 15.8)),
        ("heading2", "3.2 模式互锁与输出联动"),
        ("para", "另一组梯形图网络展示了按钮输入、模式中间位和输出线圈之间的互锁关系。可以看到，程序围绕 M0.2、M0.3、M0.4 等模式位，以及 Q1.0、Q1.1、Q1.2 等输出或状态位构建了互斥条件；同时又结合 I0.2、I0.3、I0.4 等外部输入，实现不同运行模式或不同动作段之间的切换。理论上，这类互锁逻辑可以防止多个模式同时生效，避免触摸屏按钮触发后造成控制冲突。"),
        ("figure", Figure("plc_logic_page2.png", "图7 PLC 程序中的互锁与状态切换网络", 15.8)),
        ("heading1", "4 新增 HMI 部分对原实验系统的意义"),
        ("para", "前版报告已经证明该系统能够完成运动控制任务，而本次新增的 HMI 组态与联调材料进一步说明系统已经具备了更完整的实验展示能力。对于定位跑与折返跑实验，触摸屏可以直接显示运行计数与当前输出状态；对于跨栏跑和 PID 相关实验，界面中的模拟量值和模式状态又能帮助观察变速与控制量变化趋势。也就是说，触摸屏并不替代底层控制，而是提升了实验系统在监控、验证和展示方面的层次。"),
        ("para", "从教学意义上看，加入触摸屏后，实验报告不再只停留在“PLC 程序能否驱动设备”的层面，而是扩展到了“控制系统如何被观察、如何被交互、如何被验证”的层面。这种拓展更接近完整自动化系统的工程思维。"),
        ("heading1", "5 结论"),
        ("para", "综合前版运动控制内容与本次新增截图可以认为，本实验系统已经形成了“底层 PLC 控制 + 上位 HMI 监控”的完整结构。理论层面上，通信组态、变量映射、界面设计和互锁逻辑彼此配合，使得系统不仅能够完成运动任务，还能够以较清晰的方式展示运行结果和状态变化。理论版报告的重点因此落在系统机理解释、信息流分析和结构合理性说明上。"),
        ("heading1", "参考文献"),
        ("ref", "［1］西门子（中国）有限公司. S7-200 可编程控制器系统手册[Z]."),
        ("ref", "［2］Siemens AG. TP177A/TP177B/OP177B Operating Instructions[Z]."),
        ("ref", "［3］西门子 WinCC flexible 组态软件相关资料[Z]."),
        ("ref", "［4］FRS520SE 变频器说明书[Z]."),
    ]


def practice_sections() -> list[tuple[str, object]]:
    return [
        ("abstract", "本报告在前期数字化集成控制系统实验报告的基础上，侧重整理 TP177B 触摸屏的实际组态、变量建立、界面绘制以及与 PLC 联调下载过程，形成一份偏实操过程的实验报告。报告按照“建立连接—定义变量—绘制界面—编写联动程序—下载测试”的顺序展开，结合 WinCC flexible 和 STEP7-Micro/WIN 的真实截图，说明本次实验中触摸屏监控层是如何逐步搭建完成的。结果表明，触摸屏已经能够完成数值显示、状态反馈和操作切换，并与原有运动控制程序实现联动，为系统测试与展示提供了更加直接的实验手段。"),
        ("keywords", "关键词：WinCC flexible；TP177B；组态连接；变量表；PLC 联调"),
        ("heading1", "1 实验目的与环境"),
        ("para", "本次补充实验的目标是在原有定位跑、折返跑、跨栏跑和 PID 控制程序基础上，完成触摸屏组态与 PLC 联调。具体任务包括：建立 TP177B 与 S7-200 PLC 的以太网连接；在 HMI 中创建与 PLC 地址对应的变量；设计运行监控界面；在 PLC 程序中完成通信初始化与模式互锁；最终通过下载测试验证界面显示与程序运行的一致性。"),
        ("para", "软件环境主要包括 WinCC flexible Standard 与 STEP7-Micro/WIN，硬件对象包括 TP177B 6'' color PN/DP、S7-200 PLC、变频器和实验台电机装置。前版报告中的硬件与软件总体结构继续有效，因此本次更关注组态实施与联调过程。"),
        ("figure", Figure("hardware_topology.png", "图1 原实验系统的硬件拓扑图", 15.0)),
        ("heading1", "2 触摸屏组态实施过程"),
        ("heading2", "2.1 建立 HMI 与 PLC 连接"),
        ("para", "首先在 WinCC flexible 中新建项目，选择 TP177B 6'' color PN/DP 作为 HMI 设备，并在“连接”页面建立名为“连接_1”的通信连接。截图显示，驱动程序选择为 SIMATIC S7 200，接口方式选择以太网，HMI 地址设置为 192.168.1.2，PLC 地址设置为 192.168.1.4。完成后即可形成 HMI 与 PLC 之间的基本通信通道。"),
        ("para", "在实际操作中，这一步的关键是保证 HMI 和 PLC 位于同一网段，并且地址不冲突。如果地址配置错误，后续变量刷新就无法正常进行，因此连接页面是整个组态流程的起点。"),
        ("figure", Figure("hmi_connection.png", "图2 HMI 与 PLC 的连接组态页面", 15.8)),
        ("heading2", "2.2 创建变量并绑定 PLC 地址"),
        ("para", "连接建立完成后，在“变量”页面中逐一创建需要显示或调用的变量。从截图可见，已经配置了 DWord 类型的 VD22、Int 类型的 VW20，以及多项 Bool 类型变量，分别对应 M0.2、M0.3、M0.4 和 Q1.0、Q1.1 等地址。采集周期统一设为 1 s，说明本次组态以实验观察与状态显示为主要目的。"),
        ("para", "变量创建时需要特别注意两点：一是数据类型必须与 PLC 内部寄存器或位地址匹配；二是显示用途与变量含义要提前规划清楚。例如，VD22 更适合用于高速计数值显示，VW20 适合绑定模拟量或中间值，而 M 位与 Q 位适合映射成状态灯、按钮触发或模式选择逻辑。"),
        ("figure", Figure("hmi_variables.png", "图3 HMI 变量表与地址绑定结果", 15.8)),
        ("heading2", "2.3 绘制运行监控界面"),
        ("para", "在画面设计阶段，建立了“运行测试程序”主界面，并布置了三个功能按钮区、高速计数器值显示框、模拟量值显示框、启动输入控件以及“主界面”返回按钮。通过这种布局，可以在一个页面上完成启动控制、模式观察和数据查看，减少操作时的页面切换。"),
        ("para", "从实操角度看，画面设计不仅追求美观，更重要的是便于测试。把计数值、模拟量值和按钮区放在同一页，能够使实验者在运行过程中边操作边观察，从而更快发现变量刷新异常、按钮互锁异常或者程序状态切换异常。"),
        ("figure", Figure("hmi_runtime.png", "图4 触摸屏运行界面设计结果", 12.5)),
        ("heading1", "3 PLC 联调与程序验证"),
        ("heading2", "3.1 通信初始化程序下载"),
        ("para", "在 STEP7-Micro/WIN 中下载 PLC 程序后，可以看到程序中包含 ETH1_CTRL 通信控制模块以及若干状态切换网络。第一组截图显示，程序在前部网络中通过系统位 SM0.0 触发通信初始化，为触摸屏数据访问提供前提条件。该步骤的意义在于先保证通信功能可用，再进行后续按钮联动和状态切换测试。"),
        ("figure", Figure("plc_logic_page1.png", "图5 PLC 联调程序前段网络截图", 15.8)),
        ("heading2", "3.2 按钮互锁与运行状态验证"),
        ("para", "另一组网络截图展示了模式按钮与输出线圈的互锁关系。程序围绕 M0.2、M0.3、M0.4 等中间位，配合 I0.2、I0.3、I0.4 等输入信号，实现 Q1.0、Q1.1、Q1.2 的切换控制。实验上，这意味着不同按钮不会在同一时刻同时驱动多个模式，从而避免界面操作与底层输出之间发生冲突。"),
        ("para", "从调试流程看，应先下载程序，再观察触摸屏上的变量是否正常刷新，最后通过触摸按钮或外部输入逐项验证互锁效果。如果发现某个模式按钮按下后界面无变化，通常应回到变量表与地址绑定页面检查变量类型和连接名是否一致。"),
        ("figure", Figure("plc_logic_page2.png", "图6 PLC 程序互锁网络下载验证截图", 15.8)),
        ("heading1", "4 调试现象与问题处理"),
        ("para", "本次组态与联调过程中，最需要关注的问题有三类。第一类是通信类问题，例如 IP 地址配置不在同一网段、连接驱动选择错误或访问点配置不正确；第二类是变量类问题，例如 HMI 变量类型与 PLC 地址不匹配，导致显示异常；第三类是逻辑类问题，例如按钮互锁不充分，造成多个输出状态同时有效。通过新增截图可以看到，这三类关键环节都已经在实验中被实际处理和验证。"),
        ("para", "对于整个数字化集成控制系统而言，触摸屏的加入使原有运动控制实验的调试效率明显提高。实验者不必只依赖 PLC 软件在线监视，也可以直接通过界面观察计数与状态变化，因此更适合课堂展示和实验记录。"),
        ("heading1", "5 与原有控制任务的结合"),
        ("para", "在原有五个实验题中，触摸屏监控界面最直接的作用是为定位跑、折返跑和跨栏跑提供运行状态与计数值显示，为后续位置 PID 和双闭环 PID 提供模拟量显示与状态判断入口。换句话说，本次新增部分不是独立于原系统之外的附加模块，而是把原先偏底层的控制程序进一步做成了可视化、可操作的完整实验系统。"),
        ("heading1", "6 结论"),
        ("para", "实操版报告的重点是把组态流程和联调结果写清楚。结合连接页面、变量表、运行界面和 PLC 程序截图可以确认，本次实验已经完成 TP177B 触摸屏与 S7-200 PLC 的基本联调，并能够为原有运动控制任务提供较稳定的人机交互支持。"),
        ("heading1", "参考文献"),
        ("ref", "［1］西门子（中国）有限公司. S7-200 可编程控制器系统手册[Z]."),
        ("ref", "［2］Siemens AG. TP177A/TP177B/OP177B Operating Instructions[Z]."),
        ("ref", "［3］WinCC flexible 组态使用资料[Z]."),
        ("ref", "［4］数字化集成控制系统实验平台说明资料[Z]."),
    ]


def common_toc(section_list: list[tuple[str, object]]) -> list[str]:
    toc: list[str] = []
    for kind, value in section_list:
        if kind == "heading1":
            toc.append(str(value))
    return toc


def add_figure_docx(doc: Document, figure: Figure) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(ASSET_DIR / figure.image), width=Cm(figure.width))
    add_paragraph(doc, figure.caption, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0, line_spacing=1.0)


def build_docx(output_path: Path, report_title: str, abstract_label: str, sections: list[tuple[str, object]]) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.6)
    sec.bottom_margin = Cm(2.6)
    sec.left_margin = Cm(2.6)
    sec.right_margin = Cm(2.6)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.add_run("\n\n")
    run = title_para.add_run(report_title)
    set_run_font(run, 22, bold=True)
    add_paragraph(doc, f"班级：{STUDENT_CLASS}", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0, line_spacing=1.25)
    add_paragraph(doc, f"姓名：{STUDENT_NAME}    学号：{STUDENT_ID}", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0, line_spacing=1.25)
    add_paragraph(doc, f"课程名称：{COURSE_NAME}", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0, line_spacing=1.25)
    doc.add_page_break()

    add_paragraph(doc, "目录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0, line_spacing=1.2)
    add_paragraph(doc, abstract_label, size=12, first_line_indent=0, line_spacing=1.15)
    for item in common_toc(sections):
        add_paragraph(doc, item, size=12, first_line_indent=0, line_spacing=1.15)
    doc.add_page_break()

    add_paragraph(doc, abstract_label, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=0, line_spacing=1.2)
    for kind, value in sections:
        if kind == "abstract":
            add_paragraph(doc, str(value), size=12)
        elif kind == "keywords":
            add_paragraph(doc, str(value), size=11.5, first_line_indent=0)
        elif kind == "heading1":
            add_paragraph(doc, str(value), size=16, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=0, line_spacing=1.2)
        elif kind == "heading2":
            add_paragraph(doc, str(value), size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=0, line_spacing=1.2)
        elif kind == "para":
            add_paragraph(doc, str(value), size=12)
        elif kind == "figure":
            add_figure_docx(doc, value)
        elif kind == "ref":
            add_paragraph(doc, str(value), size=11.5, first_line_indent=0, line_spacing=1.15)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)
    doc.save(str(output_path))


def tex_preamble(title: str) -> str:
    return rf"""\documentclass[12pt,a4paper]{{ctexart}}
\usepackage[left=2.6cm,right=2.6cm,top=2.6cm,bottom=2.6cm]{{geometry}}
\usepackage{{graphicx}}
\usepackage{{setspace}}
\usepackage{{hyperref}}
\usepackage{{fancyhdr}}
\usepackage{{titlesec}}
\usepackage{{caption}}
\setCJKmainfont[BoldFont=SimHei]{{SimSun}}
\setlength{{\parindent}}{{2em}}
\setlength{{\parskip}}{{0pt}}
\setlength{{\emergencystretch}}{{3em}}
\linespread{{1.28}}
\graphicspath{{{{{tex_escape(str(ASSET_DIR).replace("\\", "/"))}/}}}}
\hypersetup{{colorlinks=true,linkcolor=black,urlcolor=blue}}
\pagestyle{{fancy}}
\fancyhf{{}}
\fancyfoot[C]{{\thepage}}
\renewcommand{{\headrulewidth}}{{0pt}}
\captionsetup{{font=small}}
\titleformat{{\section}}{{\zihao{{4}}\bfseries}}{{\thesection}}{{0.6em}}{{}}
\titleformat{{\subsection}}{{\zihao{{-4}}\bfseries}}{{\thesubsection}}{{0.6em}}{{}}
\begin{{document}}
\begin{{titlepage}}
\centering
\vspace*{{2.2cm}}
{{\zihao{{2}}\bfseries {tex_escape(title)}\par}}
\vspace{{2.6cm}}
{{\zihao{{4}}班级：{tex_escape(STUDENT_CLASS)} \quad 姓名：{tex_escape(STUDENT_NAME)} \quad 学号：{tex_escape(STUDENT_ID)}\par}}
\vspace{{1.1cm}}
{{\zihao{{4}}课程名称：{tex_escape(COURSE_NAME)}\par}}
\vfill
{{\zihao{{-4}}\today\par}}
\end{{titlepage}}
\tableofcontents
\clearpage
"""


def tex_sections(abstract_heading: str, sections: list[tuple[str, object]]) -> str:
    out: list[str] = [rf"\section*{{{tex_escape(abstract_heading)}}}", rf"\addcontentsline{{toc}}{{section}}{{{tex_escape(abstract_heading)}}}"]
    refs_open = False
    for kind, value in sections:
        if kind == "abstract":
            out.append(tex_escape(str(value)))
        elif kind == "keywords":
            out.append(r"\noindent\textbf{" + tex_escape(str(value).split("：", 1)[0] + "：") + "}" + tex_escape(str(value).split("：", 1)[1]))
            out.append(r"\clearpage")
        elif kind == "heading1":
            if refs_open:
                out.append(r"\end{enumerate}")
                refs_open = False
            if str(value) == "参考文献":
                out.append(r"\section*{参考文献}")
                out.append(r"\addcontentsline{toc}{section}{参考文献}")
            else:
                out.append(rf"\section*{{{tex_escape(str(value))}}}")
                out.append(rf"\addcontentsline{{toc}}{{section}}{{{tex_escape(str(value))}}}")
        elif kind == "heading2":
            if refs_open:
                out.append(r"\end{enumerate}")
                refs_open = False
            out.append(rf"\subsection*{{{tex_escape(str(value))}}}")
        elif kind == "para":
            out.append(tex_escape(str(value)))
        elif kind == "figure":
            fig: Figure = value
            out.append(r"\begin{figure}[htbp]")
            out.append(r"\centering")
            out.append(rf"\includegraphics[width={fig.width/16.5:.2f}\textwidth]{{{fig.image}}}")
            out.append(rf"\caption{{{tex_escape(fig.caption)}}}")
            out.append(r"\end{figure}")
        elif kind == "ref":
            if not refs_open:
                out.append(r"\begin{enumerate}")
                refs_open = True
            cleaned = str(value).replace("［", "").replace("］", "")
            cleaned = cleaned.replace("1］", "").replace("2］", "").replace("3］", "").replace("4］", "")
            out.append(rf"\item {tex_escape(cleaned)}")
    if refs_open:
        out.append(r"\end{enumerate}")
    out.append(r"\end{document}")
    return "\n".join(out)


def compile_pdf(tex_path: Path) -> None:
    for _ in range(2):
        subprocess.run(
            ["xelatex", "-interaction=nonstopmode", "-halt-on-error", tex_path.name],
            cwd=tex_path.parent,
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
        )


def build_tex_pdf(build_name: str, report_title: str, abstract_heading: str, sections: list[tuple[str, object]], final_pdf_name: str) -> None:
    tex_path = BUILD_DIR / f"{build_name}.tex"
    tex_text = tex_preamble(report_title) + tex_sections(abstract_heading, sections)
    tex_path.write_text(tex_text, encoding="utf-8")
    compile_pdf(tex_path)
    shutil.copyfile(BUILD_DIR / f"{build_name}.pdf", ROOT / final_pdf_name)


def main() -> None:
    copy_assets()

    theory_title = "数字化集成控制系统实验报告（理论版）"
    practice_title = "数字化集成控制系统实验报告（实操版）"
    abstract_heading = "摘要与关键词"

    theory = theory_sections()
    practice = practice_sections()

    build_docx(ROOT / f"{theory_title}.docx", theory_title, abstract_heading, theory)
    build_docx(ROOT / f"{practice_title}.docx", practice_title, abstract_heading, practice)

    build_tex_pdf("theory_report", theory_title, abstract_heading, theory, f"{theory_title}.pdf")
    build_tex_pdf("practice_report", practice_title, abstract_heading, practice, f"{practice_title}.pdf")


if __name__ == "__main__":
    main()
