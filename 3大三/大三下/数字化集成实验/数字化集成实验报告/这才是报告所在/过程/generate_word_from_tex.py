from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
TEX_PATH = ROOT / "report.tex"
OUTPUT_PATH = ROOT / "数字化集成控制系统期末实验报告.docx"


def set_run_font(run, size: float, *, bold: bool = False, font_name: str = "宋体") -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)


def add_paragraph(
    doc: Document,
    text: str,
    *,
    size: float = 12,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent: float = 21,
    line_spacing: float = 1.35,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Pt(first_line_indent) if first_line_indent else Pt(0)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size, bold=bold)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def clean_tex_text(text: str) -> str:
    text = text.replace(r"\textit{", "")
    text = text.replace(r"\par", "")
    text = re.sub(r"\\url\{([^}]+)\}", r"\1", text)
    text = text.replace("{", "").replace("}", "")
    text = text.replace(r"\quad", " ")
    text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?", "", text)
    return " ".join(text.split()).strip()


def parse_tex() -> list[tuple[str, str]]:
    content = TEX_PATH.read_text(encoding="utf-8")
    lines = content.splitlines()
    items: list[tuple[str, str]] = []
    in_refs = False
    in_enum = False
    in_figure = False
    started = False
    current_para: list[str] = []

    def flush_para():
        nonlocal current_para
        if current_para:
            text = clean_tex_text(" ".join(current_para))
            if text:
                items.append(("para", text))
        current_para = []

    for raw in lines:
        line = raw.strip()
        if not line or line.startswith("%"):
            flush_para()
            continue
        if line.startswith(r"\begin{figure"):
            if not started:
                continue
            flush_para()
            in_figure = True
            continue
        if in_figure:
            if r"\includegraphics" in line:
                match = re.search(r"\{([^}]+)\}", line)
                if match:
                    items.append(("image", match.group(1)))
            elif r"\caption{" in line:
                cap = re.search(r"\\caption\{([^}]+)\}", line)
                if cap:
                    items.append(("caption", cap.group(1)))
            elif line.startswith(r"\end{figure"):
                in_figure = False
            continue
        if line.startswith(r"\begin{enumerate"):
            flush_para()
            in_refs = True
            in_enum = True
            continue
        if line.startswith(r"\end{enumerate"):
            flush_para()
            in_enum = False
            in_refs = False
            continue
        if in_enum and line.startswith(r"\item"):
            text = clean_tex_text(line[len(r"\item"):].strip())
            items.append(("ref", text))
            continue
        if line.startswith(r"\section*{摘要与关键词}"):
            flush_para()
            started = True
            items.append(("heading1", "摘要与关键词"))
            continue
        if not started:
            continue
        for kind, prefix in [("heading1", r"\section{"), ("heading2", r"\subsection{"), ("heading3", r"\subsubsection{")]:
            if line.startswith(prefix):
                flush_para()
                title = re.search(r"\{([^}]*)\}", line).group(1)
                items.append((kind, title))
                break
        else:
            if line.startswith(r"\appendix"):
                flush_para()
                continue
            if line.startswith(r"\section*{参考文献}"):
                flush_para()
                items.append(("heading1", "参考文献"))
                continue
            if line.startswith(r"\VerbatimInput"):
                flush_para()
                match = re.search(r"\{([^}]+)\}", line)
                if match:
                    items.append(("code", match.group(1)))
                continue
            if line.startswith(r"\titlepage") or line.startswith(r"\begin{titlepage}") or line.startswith(r"\end{titlepage}") or line.startswith(r"\tableofcontents") or line.startswith(r"\clearpage") or line.startswith(r"\addcontentsline") or line.startswith(r"\begin{document}") or line.startswith(r"\end{document}") or line.startswith(r"\small") or line.startswith(r"\normalsize") or line.startswith(r"\begin{sloppypar}") or line.startswith(r"\end{sloppypar}"):
                flush_para()
                continue
            current_para.append(line)
    flush_para()
    return items


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.6)
    section.bottom_margin = Cm(2.6)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.add_run("\n\n")
    run = p.add_run("《数字化集成控制系统》期末实验报告")
    set_run_font(run, 22, bold=True)

    add_paragraph(doc, "班级：强基班", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0, line_spacing=1.3)
    add_paragraph(doc, "姓名：匿名    学号：", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0, line_spacing=1.3)
    add_paragraph(doc, "课程名称：数字化集成控制系统", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0, line_spacing=1.3)
    add_paragraph(doc, "报告内容：定位跑、折返跑、跨栏跑、位置 PID、双闭环 PID", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0, line_spacing=1.3)
    doc.add_page_break()

    add_paragraph(doc, "目录", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0, line_spacing=1.2)
    toc_lines = [
        "摘要与关键词",
        "1 系统总体方案设计",
        "2 基础控制原理与概念",
        "3 控制方案设计",
        "4 总结与展望",
        "参考文献",
        "附录A 定位跑 STL 程序",
        "附录B 折返跑 STL 程序",
        "附录C 跨栏跑 STL 程序",
        "附录D 位置 PID STL 程序",
        "附录E 双闭环 PID STL 程序",
    ]
    for line in toc_lines:
        add_paragraph(doc, line, size=12, first_line_indent=0, line_spacing=1.2)
    doc.add_page_break()

    for kind, value in parse_tex():
        if kind == "heading1":
            add_paragraph(doc, value, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=0, line_spacing=1.2)
        elif kind == "heading2":
            add_paragraph(doc, value, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=0, line_spacing=1.2)
        elif kind == "heading3":
            add_paragraph(doc, value, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=0, line_spacing=1.2)
        elif kind == "para":
            add_paragraph(doc, value, size=12, line_spacing=1.35)
        elif kind == "image":
            img_path = ROOT / value
            if not img_path.exists():
                img_path = ROOT / "latex_assets" / value
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(img_path), width=Cm(15.8))
        elif kind == "caption":
            add_paragraph(doc, value, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=0, line_spacing=1.0)
        elif kind == "ref":
            add_paragraph(doc, value, size=11, first_line_indent=0, line_spacing=1.2)
        elif kind == "code":
            code_path = ROOT / value
            title_map = {
                "dingweipao.txt": "附录A 定位跑 STL 程序",
                "zhefanpao.txt": "附录B 折返跑 STL 程序",
                "kualanpao.txt": "附录C 跨栏跑 STL 程序",
                "position_pid.txt": "附录D 位置 PID STL 程序",
                "dual_pid.txt": "附录E 双闭环 PID STL 程序",
            }
            add_paragraph(doc, title_map.get(code_path.name, code_path.name), size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=0, line_spacing=1.2)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(8)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(code_path.read_text(encoding="utf-8"))
            set_run_font(run, 8.5, font_name="Courier New")
            doc.add_page_break()

    footer_para = doc.sections[0].footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer_para)
    doc.save(str(OUTPUT_PATH))


if __name__ == "__main__":
    build_docx()
